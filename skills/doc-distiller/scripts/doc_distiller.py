#!/usr/bin/env python3
"""
MEGA-SKILL DOCUMENT DISTILLER — V1 (Universal Cross-Platform)
Architettura modulare: L'Occhio, Il Chirurgo, L'Equalizzatore, L'Impaginatore + Il Distillatore
Portabile su Windows, macOS, Linux.
Integrato con le API di fallback LLM e scrittura note del Secondo Cervello.
"""

import os
import sys
import re
import time
import requests
from pathlib import Path

# ─────────── CONFIGURAZIONE AGNOSTICA ───────────
OLLAMA_HOST = os.environ.get("OLLAMA_HOST", "http://127.0.0.1:11434")
DEFAULT_MODEL = "deepseek-v4-flash:cloud"

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

# Tenta l'integrazione con il Secondo Cervello
try:
    # Aggiunge la cartella radice del Secondo Cervello a sys.path
    root_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".."))
    if root_path not in sys.path:
        sys.path.append(root_path)
        
    from engine.tools.vault_tools import get_vault_path, write_wiki_page, append_to_log
    from engine.utils.llm_fallback import call_llm_with_fallback
    from engine.utils.markdown import load_settings
    from google.antigravity import LocalAgentConfig
    
    SECOND_BRAIN_AVAILABLE = True
except ImportError:
    SECOND_BRAIN_AVAILABLE = False

# ─────────── UTILITY DI FALLBACK LLM ───────────
async def _call_llm_async(prompt: str, system_instructions: str):
    if SECOND_BRAIN_AVAILABLE:
        try:
            vault_path = get_vault_path()
            settings = load_settings(vault_path)
            model = settings.get("models", {}).get("ingest_agent", "gemini-3.5-flash")
            
            # Configura Vertex AI se abilitato
            auth = settings.get("google_auth", {})
            kwargs = {}
            if auth.get("use_vertex", False):
                kwargs["vertex"] = True
                if auth.get("project_id"):
                    kwargs["project"] = auth["project_id"]
                if auth.get("location"):
                    kwargs["location"] = auth["location"]
                    
            config = LocalAgentConfig(
                model=model,
                system_instructions=system_instructions,
                **kwargs
            )
            print(f"  [Doc Distiller] Utilizzo LLM del Secondo Cervello ({model})...")
            return await call_llm_with_fallback(prompt, system_instructions, config)
        except Exception as e:
            print(f"  [Doc Distiller - Avviso] Chiamata Secondo Cervello fallita ({e}). Tento fallback locale su Ollama...")
            
    # Chiamata di fallback nativa su Ollama
    payload = {
        "model": DEFAULT_MODEL,
        "messages": [
            {"role": "system", "content": system_instructions},
            {"role": "user", "content": prompt}
        ],
        "options": {"temperature": 0.2, "num_ctx": 32768},
        "stream": False
    }
    resp = requests.post(f"{OLLAMA_HOST}/api/chat", json=payload, timeout=240)
    return resp.json().get("message", {}).get("content", "")

def call_llm_sync(prompt: str, system_instructions: str) -> str:
    import asyncio
    try:
        loop = asyncio.get_event_loop()
    except RuntimeError:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        
    if loop.is_running():
        # Esecuzione in thread pool se c'è un loop in corso (es. in FastAPI)
        import concurrent.futures
        with concurrent.futures.ThreadPoolExecutor() as executor:
            future = executor.submit(lambda: asyncio.run(_call_llm_async(prompt, system_instructions)))
            return future.result()
    else:
        return loop.run_until_complete(_call_llm_async(prompt, system_instructions))

# ─────────── MICRO-SKILL 1: L'OCCHIO ───────────
def estrai_testo_universale(file_path):
    fp = Path(file_path)
    ext = fp.suffix.lower()
    print(f"[L'Occhio] Estrazione in corso su {fp.name}...")
    
    if ext == ".pdf":
        return _estrai_pdf(fp)
    elif ext in {".txt", ".md", ".markdown"}:
        return fp.read_text(encoding="utf-8", errors="replace"), 1
    elif ext == ".docx":
        return _estrai_docx(fp)
    else:
        try:
            return fp.read_text(encoding="utf-8", errors="replace"), 1
        except Exception:
            raise RuntimeError(f"Formato non supportato nativamente: {ext}")

def _estrai_pdf(fp):
    testo = []
    num_pages = 0
    if pdfplumber:
        try:
            with pdfplumber.open(str(fp)) as pdf:
                num_pages = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    t = page.extract_text()
                    if t:
                        testo.append(f"\n\n[PAGINA {i + 1}]\n\n{t}")
            return "".join(testo), num_pages
        except Exception as e:
            print(f"  [Warning] pdfplumber fallito, provo fallback su pypdf: {e}")

    if PdfReader:
        reader = PdfReader(str(fp))
        num_pages = len(reader.pages)
        for i, page in enumerate(reader.pages):
            t = page.extract_text()
            if t:
                testo.append(f"\n\n[PAGINA {i + 1}]\n\n{t}")
        return "".join(testo), num_pages
    
    raise ImportError("Installa 'pdfplumber' o 'pypdf' per estrarre dati da PDF.")

def _estrai_docx(fp):
    try:
        import docx
        doc = docx.Document(str(fp))
        return "\n".join([p.text for p in doc.paragraphs if p.text]), 1
    except ImportError:
        import zipfile
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(str(fp)) as zf:
            xml_bytes = zf.read("word/document.xml")
        root = ET.fromstring(xml_bytes)
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        parts = [n.text for n in root.iter(f"{ns}t") if n.text]
        return "\n".join(parts), 1

# ─────────── MICRO-SKILL 2: IL CHIRURGO ───────────
def chunking_adattivo(testo, max_tok=24000, overlap_paragrafi=2):
    print(f"[Il Chirurgo] Segmentazione testo adattiva...")
    paragrafi = [p.strip() for p in re.split(r"\n\s*\n", testo) if p.strip()]
    chunks = []
    current, cur_tok, cid = [], 0, 0
    
    for para in paragrafi:
        pt = len(para) // 4
        if cur_tok + pt > max_tok and current:
            chunks.append({"id": cid, "text": "\n\n".join(current), "token_count": cur_tok})
            cid += 1
            overlap = current[-overlap_paragrafi:] if overlap_paragrafi > 0 else []
            current = list(overlap)
            cur_tok = sum(len(p) // 4 for p in current)
        current.append(para)
        cur_tok += pt
        
    if current:
        chunks.append({"id": cid, "text": "\n\n".join(current), "token_count": cur_tok})
    return chunks

# ─────────── MICRO-SKILL 3: L'EQUALIZZATORE ───────────
PROMPT_DISTILLER = """Riscrivi il testo seguente in un saggio di sintesi strutturato, denso e discorsivo.
Usa uno stile lucido ed elegante, ideale per la ricerca, la consultazione rapida e lo studio sistemico.

LINEE GUIDA:
1. TERMINI CHIAVE: Fornisci una spiegazione sintetica tra parentesi la prima volta che introduci termini specialistici o di nicchia.
2. ENFASI: Applica il **grassetto** in modo selettivo su nomi di autori fondamentali, date, eventi spartiacque o definizioni cardine.
3. FLUIDITÀ: Evita elenchi puntati o numerati nel corpo del testo. Mantieni una prosa discorsiva ma rigorosa.
4. METADATI DI FONTE: Concludi ogni blocco concettuale con la notazione [Fonte: Blocco {chunk_id}].

Alla fine del riassunto inserisci tassativamente queste due sezioni formattate esattamente così:
### Scheda dei Concetti Fondamentali
Concetto: [Termine] - [Spiegazione estesa e approfondita]

### Domande e Problemi di Verifica
Domanda: [Domanda analitica] - [Risposta esaustiva basata solo sul testo]

TESTO DA ELABORARE:
{testo}"""

def esegui_mappatura_llm(chunk, chunk_id, lang):
    system_instructions = "Sei un esperto di analisi documentale ed estrazione della conoscenza."
    prompt = PROMPT_DISTILLER.format(chunk_id=chunk_id+1, testo=chunk["text"])
    if lang == "en":
        prompt += "\nTranslate the entire final response into fluent English."
    return call_llm_sync(prompt, system_instructions)

# ─────────── MICRO-SKILL 4: L'IMPAGINATORE ───────────
def genera_docx(testo_completo, output_path, titolo):
    try:
        import docx
        from docx.shared import Pt, Inches
    except ImportError:
        print("[L'Impaginatore] 'python-docx' non trovato. Salvo in Markdown standard.")
        Path(output_path).with_suffix('.md').write_text(testo_completo, encoding="utf-8")
        return

    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_after = Pt(6)

    doc.add_heading(f"Dossier di Sintesi & Distillazione", level=0)
    doc.add_paragraph(f"Documento sorgente: {titolo}\nGenerato via Antigravity Document Distiller\nData: {time.strftime('%Y-%m-%d')}")
    doc.add_page_break()

    bold_pattern = re.compile(r"\*\*(.+?)\*\*")
    for line in testo_completo.split("\n"):
        line_str = line.strip()
        if not line_str:
            continue
        
        if line_str.startswith("### "):
            doc.add_heading(line_str[4:], level=3)
        elif line_str.startswith("## "):
            doc.add_heading(line_str[3:], level=2)
        elif line_str.startswith("# "):
            doc.add_heading(line_str[2:], level=1)
        else:
            p = doc.add_paragraph()
            pieces = bold_pattern.split(line_str)
            for idx, piece in enumerate(pieces):
                run = p.add_run(piece)
                if idx % 2 == 1:
                    run.bold = True
                    
    doc.save(output_path)
    print(f"[L'Impaginatore] File esportato con successo in: {output_path}")

# ─────────── BONUS SKILL: IL DISTILLATORE ───────────
def distilla_conoscenza_interattiva(testo_sintesi, base_path, titolo_pulito):
    print(f"\n{'─'*40}\n🧠 SINTESI COMPLETATA CON SUCCESSO.")
    print("Desideri distillare questo materiale in un artefatto di conoscenza riutilizzabile?")
    print("1. Genera una SKILL OPERATIVA per LLM (Formato Antigravity/OpenCode .md)")
    print("2. Esporta Scheda Concettuale Atomica per il tuo SECOND BRAIN (Obsidian wiki page)")
    print("3. No, concludi l'operazione.")
    
    scelta = input("Seleziona un'opzione (1/2/3): ").strip()
    if scelta not in {"1", "2"}:
        print("Operazione conclusa!")
        return

    system_instructions = "Sei un ingegnere della conoscenza e scrittore tecnico esperto."
    prompt_distillazione = ""
    file_suffix = ""
    
    if scelta == "1":
        file_suffix = "_SKILL.md"
        prompt_distillazione = f"Analizza questa sintesi documentale ed estrai le logiche operative in formato manifesto .md per una Skill di un LLM (Framework Antigravity):\n{testo_sintesi}"
    elif scelta == "2":
        prompt_distillazione = f"Trasforma questa sintesi in una nota concettuale approfondita per il Second Brain. Usa il formato Markdown ed evita intestazioni di primo livello duplicate. Estrai definizioni, implicazioni teoriche e connessioni ad altri concetti:\n{testo_sintesi}"

    print("\n[Il Distillatore] Generazione dell'artefatto in corso...")
    
    try:
        artefatto = call_llm_sync(prompt_distillazione, system_instructions)
        
        if scelta == "1":
            out_file = Path(base_path) / f"{titolo_pulito}{file_suffix}"
            out_file.write_text(artefatto, encoding="utf-8")
            print(f"🔥 Artefatto Skill salvato in: {out_file}")
        elif scelta == "2":
            if SECOND_BRAIN_AVAILABLE:
                # Salva direttamente nel vault come nota concetto!
                concept_filename = f"wiki/concepts/General/{titolo_pulito}.md"
                write_wiki_page(
                    concept_filename, 
                    body=f"# Concetto: {titolo_pulito.replace('_', ' ')}\n\n{artefatto}", 
                    frontmatter={"type": "concept", "tags": ["distillato", "sintesi"]}
                )
                vault_path = get_vault_path()
                print(f"🔥 Nota concettuale salvata direttamente nel vault Obsidian: {os.path.join(vault_path, concept_filename)}")
                # Auto commit se abilitato nel modulo git_ops
                try:
                    from engine.git_ops import auto_commit
                    auto_commit(vault_path, f"[Doc Distiller] Distillato nuovo concetto [[{titolo_pulito}]]")
                except Exception:
                    pass
            else:
                out_file = Path(base_path) / f"{titolo_pulito}_SECOND_BRAIN.md"
                out_file.write_text(artefatto, encoding="utf-8")
                print(f"🔥 Scheda salvata localmente in: {out_file}")
                
    except Exception as e:
        print(f"⚠️ Impossibile distillare l'artefatto: {e}")

# ─────────── PIPELINE ───────────
def run_pipeline(target_file, lang="it"):
    path_sorgente = Path(target_file).resolve()
    if not path_sorgente.exists():
        print(f"Errore: Il file {target_file} non esiste.", file=sys.stderr)
        sys.exit(1)
        
    base_dir = path_sorgente.parent
    titolo_clean = path_sorgente.stem.replace(" ", "_")
    
    testo_estratto, _ = estrai_testo_universale(path_sorgente)
    chunks = chunking_adattivo(testo_estratto)
    
    print(f"[L'Equalizzatore] Elaborazione di {len(chunks)} blocchi via LLM...")
    sintesi_parziali = []
    for idx, chunk in enumerate(chunks):
        res = esegui_mappatura_llm(chunk, idx, lang)
        if res:
            sintesi_parziali.append(res)
            
    testo_sintesi_finale = "\n\n".join(sintesi_parziali)
    output_docx = base_dir / f"Distillato_{titolo_clean}.docx"
    genera_docx(testo_sintesi_finale, output_docx, path_sorgente.name)
    distilla_conoscenza_interattiva(testo_sintesi_finale, base_dir, titolo_clean)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python doc_distiller.py <percorso-documento> [lingua=it|en]")
        sys.exit(1)
    run_pipeline(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else "it")
